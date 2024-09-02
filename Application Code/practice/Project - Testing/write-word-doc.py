from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Crypto Analysis Tool', 0)

# Content
doc.add_paragraph(
    'Crypto Analysis Tool is a tool for analyzing the transaction histories of cryptocurrencies, specifically Ethereum, '
    'to identify illicit activities such as phishing, money laundering, and other suspicious behaviors. '
    'This project encompasses the development of machine learning models for anomaly detection and a web application '
    'to make these insights accessible to users.'
)

# Table of Contents
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('1. Project Overview')
doc.add_paragraph('2. Local Development')
doc.add_paragraph('   - Installing Requirements')
doc.add_paragraph('3. Machine Learning Model Training')
doc.add_paragraph('   - Phishing Detection')
doc.add_paragraph('   - Money Laundering Detection')
doc.add_paragraph('4. Web Application')
doc.add_paragraph('   - Overview')
doc.add_paragraph('   - Key Features')
doc.add_paragraph('   - Running the Web Application')
doc.add_paragraph('   - Web Application Configuration')
doc.add_paragraph('5. University of Buckingham')

# Project Overview
doc.add_heading('Project Overview', level=1)
doc.add_paragraph(
    'The Crypto Analysis Tool aims to provide a comprehensive, automated solution for detecting illicit activities within '
    'the Ethereum blockchain. The tool includes machine learning models trained on real transaction data to identify patterns '
    'indicative of fraudulent activities such as phishing and money laundering. The models are integrated into a web application '
    'that allows users to perform in-depth analysis of blockchain transactions.'
)

# Local Development
doc.add_heading('Local Development', level=1)
doc.add_paragraph(
    'To set up a local development environment, use the provided Docker Compose configuration. This setup ensures all dependencies '
    'are correctly installed and the environment mirrors the production setting.'
)
doc.add_paragraph(
    'Note: We highly recommend using Docker Compose for setting up the project, as it simplifies the process and ensures that the '
    'environment is consistent across different systems. Docker Compose automatically handles all configurations, including environment '
    'variables, database settings, and static files.'
)
doc.add_paragraph(
    'You will still need to ensure the trained models are added to the models directory'
)

doc.add_paragraph('- Phishing Detection')
doc.add_paragraph('- Money Laundering Detection')

doc.add_heading('Why Use Docker Compose?', level=2)
doc.add_paragraph(
    'Simplifies Configuration: Docker Compose manages all dependencies and configurations in a single file, eliminating the need '
    'for manual setup in settings.py.'
)
doc.add_paragraph(
    'Ensures Consistency: The Docker environment mirrors the production environment, reducing the risk of inconsistencies or errors.'
)
doc.add_paragraph(
    'Quick Setup: Set up the entire project environment with a single command.'
)
doc.add_paragraph(
    'Note: the static files may not work in the docker-compose method'
)

doc.add_heading('Steps to Start Local Development Using Docker Compose:', level=2)
doc.add_paragraph(
    '1. Clone the repository:'
)
doc.add_paragraph(
    "```powershell\ngit clone https://github.com/RuaidhriBrown/CryptoAnalysis.git\n"
    "cd 'CryptoAnalysis\\Application Code\\src\\Host Applications\\Crypto.Tracker.Web.UI\\'\n```"
)
doc.add_paragraph(
    '2. Run Docker Compose to build and start the services:'
)
doc.add_paragraph(
    "Run the following: \n```powershell\ndocker-compose up --build\n```"
)
doc.add_paragraph(
    'This command will build the Docker images and start the containers for the web application and the database. '
    'All configurations, including environment variables and database settings, are automatically managed by Docker Compose.'
)
doc.add_paragraph(
    '3. Access the Web Application:'
)
doc.add_paragraph(
    'Once the Docker containers are up and running, you can access the web application in your browser at http://localhost:8000.'
)
doc.add_paragraph(
    '4. Login to the user:'
)
doc.add_paragraph(
    '- Username: admin\n- password: X1B2#WXYZ123a'
)

doc.add_heading('Avoid Manual Configuration', level=2)
doc.add_paragraph(
    "If you choose to use Docker Compose, you don't need to manually configure the settings.py file for environment-specific settings "
    "like SECRET_KEY, DEBUG, ALLOWED_HOSTS, or DATABASES. Docker Compose handles these settings through environment variables defined in "
    "the docker-compose.yml file."
)
doc.add_paragraph(
    'However, if you still prefer manual setup, please refer to the Web Application Configuration section below.'
)

doc.add_heading('Installing Requirements', level=2)
doc.add_paragraph(
    'Before running the project, you need to install all the required Python packages for both the machine learning models and the web application.'
)

doc.add_heading('Installing Requirements for Machine Learning Model Training', level=3)
doc.add_paragraph(
    '1. Navigate to the project directory:'
)
doc.add_paragraph(
    "```powershell\ncd 'CryptoAnalysis\\Application Code\\practice\\Project - Testing\\'\n```"
)
doc.add_paragraph(
    '2. Install the Python dependencies:'
)
doc.add_paragraph(
    'Make sure you have Python and pip installed. If you\'re using a virtual environment (recommended), activate it first.'
)
doc.add_paragraph(
    'Run the following command to install the required packages for the machine learning models:'
)
doc.add_paragraph(
    "```powershell\npip install -r requirements.txt\n```"
)
doc.add_paragraph(
    '3. Installing Requirements for the Web Application:'
)
doc.add_paragraph(
    "Navigate to the web application directory: 'cd CryptoAnalysis/Application Code/src/Host Applications/Crypto.Tracker.Web.UI'"
)
doc.add_paragraph(
    "```powershell\npip install -r requirements.txt\n```"
)
doc.add_paragraph(
    'Ensure that the requirements.txt file, which contains all necessary dependencies for running the Django web application, is located in the web application directory.'
)
doc.add_paragraph(
    'Note: If you encounter any issues during installation, check that your Python version and pip are up to date. You may also need to install additional system dependencies if required by specific Python packages.'
)

# Machine Learning Model Training
doc.add_heading('Machine Learning Model Training', level=1)
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The machine learning component of the Crypto Analysis Tool consists of several models designed to detect illicit activities on the Ethereum blockchain, '
    'specifically targeting phishing and money laundering behaviors. The models use Random Forest and Graph Neural Networks (GNNs) to analyze transaction '
    'data and identify suspicious patterns.'
)

doc.add_heading('Phishing Detection', level=3)
doc.add_paragraph(
    'Phishing detection utilizes three datasets:'
)
doc.add_paragraph(
    '- Aggregated Ethereum Wallets\n'
    '- Ethereum Transaction Data\n'
    '- ERC20 Token Transaction Data'
)
doc.add_paragraph(
    'The phishing dataset is based on data from [Kaggle](https://www.kaggle.com/datasets/vagifa/ethereum-frauddetection-dataset/data). '
    'This dataset provides a list of wallets from which transactions and ERC20 transactions were collected using the Etherscan API. '
    'The resulting data was used to create an aggregated Ethereum Wallets dataset, representing known phishing accounts. Note that the transactions involve '
    'recognized phishing accounts but may not necessarily represent genuine phishing activities.'
)

doc.add_heading('Money Laundering Detection', level=3)
doc.add_paragraph(
    'Money laundering detection focuses on datasets derived from incidents like the Upbit hack, including:'
)
doc.add_paragraph(
    '- Aggregated Ethereum Wallets\n'
    '- Ethereum Transaction Data\n'
    '- ERC20 Token Transaction Data'
)
doc.add_paragraph(
    'These datasets were created using the script located at:\n'
    '`CryptoAnalysis/Application Code/practice/Project - Testing/moneyLaundering/fetch_data/getMoneyLaunderingData1Layer.py`.'
)

doc.add_heading('Developing the Upbit-Hack Dataset for Money Laundering Detection', level=4)
doc.add_paragraph(
    'The Upbit hack dataset was developed to train models for detecting money laundering activities on the Ethereum blockchain. '
    'The dataset construction involved several key steps:'
)
doc.add_paragraph(
    '1. API Configuration and Setup: The Etherscan API was used to retrieve Ethereum transaction data. To comply with API rate limits, '
    'a maximum of 5 API calls per second was enforced, and retries were handled for failed requests.'
)
doc.add_paragraph(
    '2. Data Collection: Transaction data was collected using the `get_normal_transactions` and `get_erc20_transactions` functions. '
    'These functions fetch normal Ether transactions and ERC20 token transactions associated with a given Ethereum address, filtering '
    'the results to include only relevant transactions that meet specific criteria (e.g., minimum Ether sent and a specific start time).'
)
doc.add_paragraph(
    '3. Address Analysis: Each Ethereum address (wallet) is analyzed for potential involvement in money laundering activities using '
    'the `analyze_wallet` function. Key metrics such as Fast-In Fast-Out (FIFO) ratio, small-volume transaction ratio, zero-out ratio, '
    'and network density are calculated to identify suspicious behavior.'
)
doc.add_paragraph(
    '4. Detection of Money Laundering: The `detect_money_laundering` function applies predefined thresholds to determine if a wallet is suspicious. '
    'The thresholds include FIFO ratio, small-volume ratio, zero-out ratio, network density, and minimum balance. If a wallet meets or exceeds the thresholds, '
    'it is flagged as suspicious.'
)
doc.add_paragraph(
    '5. Interacting Wallets: The `analyze_interacted_wallets` function identifies wallets that have interacted with the initial suspicious address. '
    'It recursively analyzes these wallets to uncover potential layers of money laundering activity, expanding the dataset to include a broader network of transactions.'
)
doc.add_paragraph(
    '6. Data Compilation: Data is accumulated in global dataframes for all analyzed wallets, normal transactions, and ERC20 transactions. '
    'These dataframes are then saved to CSV files (`all_wallet_summaries.csv`, `all_normal_transactions.csv`, and `all_erc20_transactions.csv`) for further analysis and model training.'
)
doc.add_paragraph(
    '7. Create "Ethereum Transaction" and "ERC20 Token Transaction" Datasets: Combine normal and ERC20 transactions from both suspicious and legitimate wallets.'
)
doc.add_paragraph(
    '8. Create "Aggregated Ethereum Wallets" Dataset: Aggregate the "Ethereum Transaction Data" and "ERC20 Token Transaction Data" to derive features.'
)

# Training the Models
doc.add_heading('Training the Models', level=2)
doc.add_paragraph(
    'To train the machine learning models locally:'
)

doc.add_heading('Phishing model code:', level=3)
doc.add_paragraph(
    'First ensure the three following datasets are present in the directory '
    '"CryptoAnalysis\\Application Code\\practice\\Project - Testing\\datasets\\ethereum\\phishing\\":'
)
doc.add_paragraph(
    '1. "combined_er20.csv"\n'
    '2. "combined_transactions.csv"\n'
    '3. "transaction_dataset_even.csv"'
)
doc.add_paragraph(
    'In the directory: "CryptoAnalysis\\Application Code\\practice\\Project - Testing\\Phishing-ML\\" there are three Python scripts:'
)
doc.add_paragraph(
    '"LR-machine learning er20.py"\n'
    '"LR-machine learning transactions.py"\n'
    '"LR-machine learning wallets v3.py"'
)
doc.add_paragraph(
    'Running these will each create a trained model and the related features PKL file in the '
    '"CryptoAnalysis\\Application Code\\practice\\Project - Testing\\results\\" directory.'
)

doc.add_heading('Money Laundering model code:', level=3)
doc.add_paragraph(
    'First ensure the three following datasets are present in the directory '
    '"CryptoAnalysis\\Application Code\\practice\\Project - Testing\\datasets\\ethereum\\MoneyLaundering\\":'
)
doc.add_paragraph(
    '1. "combined_er20_transaction_data.csv"\n'
    '2. "combined_normal_transactions_data.csv"\n'
    '3. "combined_transaction_data.csv"'
)
doc.add_paragraph(
    'In the directory: "CryptoAnalysis\\Application Code\\practice\\Project - Testing\\moneyLaundering\\" there are three Python scripts:'
)
doc.add_paragraph(
    '"LR-machine learning er20.py"\n'
    '"LR-machine learning transactions.py"\n'
    '"LR-machine learning wallets.py"'
)
doc.add_paragraph(
    'Running these will each create a trained model and the related features PKL file in the '
    '"CryptoAnalysis\\Application Code\\practice\\Project - Testing\\results\\" directory.'
)

# Web Application Section
doc.add_heading('Web Application', level=1)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The web application is developed using Django and serves as the user interface for interacting with the trained models and '
    'analyzing Ethereum blockchain data. It allows users to upload transaction data, run model analyses, and view detailed visualizations of the results.'
)

doc.add_heading('Key Features', level=2)
doc.add_paragraph(
    '- **Data Upload and Analysis**: Users can download Ethereum transaction data from etherscan and run analyses using trained machine learning models.\n'
    '- **Visualizations and Reporting**: The app provides graphical representations of transaction patterns and allows for exporting analysis results.\n'
    '- **User Interface web app**: Facilitates easy navigation, model execution, and result interpretation, supporting digital forensic investigations.'
)

doc.add_heading('Running the Web Application', level=2)
doc.add_paragraph(
    'Required:\n'
    '6 models trained in the previous section saved into the respective directories:'
)

doc.add_paragraph(
    '"CryptoAnalysis\\Application Code\\src\\Host Applications\\Crypto.Tracker.Web.UI\\models\\ethereum\\MoneyLaundering" have the following models:'
)
doc.add_paragraph(
    '1. "random_forest_moneyLaundering_detector_aggregated.pkl"\n'
    '2. "random_forest_moneyLaundering_detector_aggregated_features.pkl"\n'
    '3. "random_forest_moneyLaundering_detector_er20.pkl"\n'
    '4. "random_forest_moneyLaundering_detector_er20_features.pkl"\n'
    '5. "random_forest_moneyLaundering_detector_transactions.pkl"\n'
    '6. "random_forest_moneyLaundering_detector_transactions_features.pkl"'
)
doc.add_paragraph(
    'and\n'
    '"CryptoAnalysis\\Application Code\\src\\Host Applications\\Crypto.Tracker.Web.UI\\models\\ethereum\\phishing" have the following files:'
)
doc.add_paragraph(
    '1. "random_forest_phishing_detector_aggregated.pkl"\n'
    '2. "random_forest_phishing_detector_aggregated_features.pkl"\n'
    '3. "random_forest_phishing_detector_er20.pkl"\n'
    '4. "random_forest_phishing_detector_er20_features.pkl"\n'
    '5. "random_forest_phishing_detector_transactions.pkl"\n'
    '6. "random_forest_phishing_detector_transactions_features.pkl"'
)

# Web Application Configuration
doc.add_heading('Web Application Configuration', level=2)
doc.add_paragraph(
    'To properly set up the Django web application for the Crypto Analysis Tool, users must adjust several settings in the `settings.py` file and configure their environment accordingly.'
)

doc.add_heading('Required Configuration Steps', level=3)
doc.add_paragraph('1. **Secret Key**')
doc.add_paragraph(
    '- **Location in `settings.py`**:\n'
    '```python\nSECRET_KEY = \'b1be7575-abe1-4fe1-9c95-dce1e276f171\'\n```'
)
doc.add_paragraph(
    '- **What to do**:\n'
    'Set the `ETHERSCAN_API_KEY` environment variable in your shell or environment manager:\n'
    '- **For PowerShell**:\n'
    '```powershell\n$env:ETHERSCAN_API_KEY = "<Your_API_Key>"\n```\n'
    '- **For Linux/macOS Bash**:\n'
    '```bash\nexport ETHERSCAN_API_KEY="<Your_API_Key>"\n```\n'
    '- **For Windows Command Prompt**:\n'
    '```cmd\nset ETHERSCAN_API_KEY=<Your_API_Key>\n```\n'
)

doc.add_paragraph('2. **Debug Mode**')
doc.add_paragraph(
    '- **Location in `settings.py`**:\n'
    '```python\nDEBUG = True\n```'
)
doc.add_paragraph(
    '- **What to do**:\n'
    'Set `DEBUG = False` for production environments to prevent sensitive data from being exposed. In a development environment, keep it set to `True` to enable detailed error messages.'
)

doc.add_paragraph('3. **Allowed Hosts**')
doc.add_paragraph(
    '- **Location in `settings.py`**:\n'
    '```python\nALLOWED_HOSTS = []\n```'
)
doc.add_paragraph(
    '- **What to do**:\n'
    'Add the domain names or IP addresses of your deployment environment to the `ALLOWED_HOSTS` list to allow Django to serve your app on those hosts. For example:\n'
    '```python\nALLOWED_HOSTS = [\'localhost\', \'127.0.0.1\', \'yourdomain.com\']\n```'
)

doc.add_paragraph('4. **Database Configuration**')
doc.add_paragraph(
    '- **Location in `settings.py`**:\n'
    '```python\n'
    'DATABASES = {\n'
    '    \'default\': {\n'
    '        \'ENGINE\': \'django.db.backends.postgresql\',\n'
    '        \'NAME\': \'CryptoTracker\',\n'
    '        \'USER\': \'crypto-postgres\',\n'
    '        \'PASSWORD\': \'X1B2#WXYZ123a\',\n'
    '        \'HOST\': \'localhost\',\n'
    '        \'PORT\': \'5432\',\n'
    '        \'OPTIONS\': {\n'
    '            \'options\': \'-c search_path=dev_crypto_tracker,public\'\n'
    '        }\n'
    '    }\n'
    '}\n```'
)
doc.add_paragraph(
    '- **What to do**:\n'
    'Ensure that PostgreSQL is installed and configured correctly on your system. You may need to modify the following database settings to match your environment:\n'
    '- `NAME`: Your PostgreSQL database name.\n'
    '- `USER`: Your PostgreSQL username.\n'
    '- `PASSWORD`: Your PostgreSQL password.\n'
    '- `HOST`: The address of your PostgreSQL server (use `localhost` for local development).\n'
    '- `PORT`: The port number for your PostgreSQL server (default is `5432`).'
)

doc.add_paragraph('5. **Static and Media Files**')
doc.add_paragraph(
    '- **Location in `settings.py`**:\n'
    '```python\n'
    'STATIC_URL = \'/static/\'\n'
    'STATIC_ROOT = posixpath.join(*(BASE_DIR.split(os.path.sep) + [\'static\']))\n'
    'MEDIA_URL = \'/media/\'\n'
    'MEDIA_ROOT = os.path.join(BASE_DIR, \'media\')\n```'
)
doc.add_paragraph(
    '- **What to do**:\n'
    'Ensure `STATIC_ROOT` and `MEDIA_ROOT` directories are correctly set for your deployment environment. These should exist and be writable by the web server.\n'
    'In production, run `python manage.py collectstatic` to gather all static files into the `STATIC_ROOT` directory.'
)

doc.add_paragraph('6. **Environment Variables**')
doc.add_paragraph(
    '- **ETHERSCAN API Key**:\n'
    'The application requires an Etherscan API key to access blockchain data. This key should be stored as an environment variable.\n'
    '- **Location in `settings.py`**:\n'
    '```python\nETHERSCAN_API_KEY = os.getenv("ETHERSCAN_API_KEY", "your-default-api-key")\n```'
)
doc.add_paragraph(
    '- **What to do**:\n'
    'Set the `ETHERSCAN_API_KEY` environment variable in your shell or environment manager:\n'
    '- **For PowerShell**:\n'
    '```powershell\n$env:ETHERSCAN_API_KEY = "<Your_API_Key>"\n```\n'
    '- **For Linux/macOS Bash**:\n'
    '```bash\nexport ETHERSCAN_API_KEY="<Your_API_Key>"\n```\n'
    '- **For Windows Command Prompt**:\n'
    '```cmd\nset ETHERSCAN_API_KEY=<Your_API_Key>\n```\n'
)

doc.add_paragraph('7. **Middleware and Installed Apps**')
doc.add_paragraph(
    '- **Middleware**: The default middleware is set up for standard security and session management. Customize further as needed for your environment.\n'
    '- **Installed Apps**: Ensure all necessary apps, including Django default apps and `webview`, are listed under `INSTALLED_APPS` in `settings.py`.'
)

doc.add_paragraph('8. **Timezone and Localization**')
doc.add_paragraph(
    '- **Location in `settings.py`**:\n'
    '```python\nLANGUAGE_CODE = \'en-us\'\nTIME_ZONE = \'UTC\'\n```'
)
doc.add_paragraph(
    '- **What to do**: Adjust `LANGUAGE_CODE` and `TIME_ZONE` according to your localization needs.'
)

doc.add_paragraph(
    'After completing these configurations, your Django application should be properly set up to run in your local or production environment.'
)
doc.add_paragraph(
    'The web application will be accessible at `http://localhost:8000`.'
)

doc.add_heading('Creating an Admin User', level=3)
doc.add_paragraph(
    'To manage your Django application through the admin interface, you need to create an admin user. Follow these steps using Visual Studio:'
)
doc.add_paragraph(
    '1. **Open the Django Project in Visual Studio**:\n'
    '- Ensure you have your Django project open in Visual Studio.'
)
doc.add_paragraph(
    '2. **Right-click on the Django Project** in the Solution Explorer pane.'
)
doc.add_paragraph(
    '3. **Navigate to Python > Django**:\n'
    '- Hover over **Python** to expand the menu.\n'
    '- Then hover over **Django** to see more options.'
)
doc.add_paragraph(
    '4. **Select \'Create Superuser\'**:\n'
    '- Click on **Create Superuser** from the dropdown menu.'
)
doc.add_paragraph(
    '5. **Follow the Prompts in the Output Window**:\n'
    '- Visual Studio will prompt you in the Output window to enter the details for your superuser.\n'
    '- Provide a username, email, and password when prompted.'
)
doc.add_paragraph(
    '6. **Access the Django Admin Interface**:\n'
    '- Open your web browser and go to `http://localhost:8000/admin`.\n'
    '- Log in using the admin credentials you just created.'
)

# University of Buckingham
doc.add_heading('University of Buckingham', level=1)
doc.add_paragraph(
    "This repository was developed as part of Ruaidhri Brown's MSc in Innovative Computing at the University of Buckingham. "
    "The project addresses critical challenges in digital forensics and cybersecurity by providing tools to detect and analyze "
    "illicit activities on the Ethereum blockchain, thereby aiding law enforcement and research in blockchain security."
)

# Save the document
doc.save('Crypto_Analysis_Tool.docx')
